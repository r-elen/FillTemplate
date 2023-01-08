import re
from docx import Document
import docx2txt as d
from datetime import datetime


SERVICES = ["Анализ нормативно-правовой базы и формирование правовой позиции", "Составление проекта",
             "3", "4", "5", "6",]


def create_num(number: str) -> str:
    d = datetime.today().strftime('%d%m%y')
    return ''.join([d, number])


def create_date():
    dd = datetime.today().strftime('%d')
    yyyy = datetime.today().strftime('%Y')

    month = datetime.today().strftime('%m')
    dic = {'01': "января", '02': "февраля", '03': "марта", '04': "апреля", '05': "мая", '06': "июня", '07': "июля", '08': "августа", '09': "сентября", '10': "октября", '11': "ноября", '12': "декабря"}
    mm = dic[month]

    return ' '.join([dd, mm, yyyy])


def create_full_name(last_name: str, name: str, middle_name=None):
    if len(middle_name) == 0:
        full_name = ' '.join([last_name, name])
    else:
        full_name = ' '.join([last_name, name, middle_name])
    return full_name


def create_initials(last_name: str, name: str, middle_name=None):
    """
    :param last_name: last name
    :param name: first name
    :param middle_name: middle name
    :return: None
    """
    if len(middle_name) == 0:
        name_initials = '.'.join([name[0], last_name])
    else:
        name_initials = '.'.join([name[0], middle_name[0], last_name])
    return name_initials


def docx_replace_regex(doc_obj, regex, replace):

    for table in doc_obj.tables:
        for row in table.rows:
            for cell in row.cells:
                docx_replace_regex(cell, regex, replace)

    for p in doc_obj.paragraphs:
        if regex.search(p.text):
            inline = p.runs
            # Loop added to work with runs (strings with same style)

            # The text to be replaced can be split over several runs so search through,
            # identify which runs need to have text replaced then replace the text in those identified
            started = False
            key_index = 0
            # found_runs is a list of (inline index, index of match, length of match)
            found_runs = list()
            found_all = False
            replace_done = False

            for i in range(len(inline)):

                # case 1: found in single run so short circuit the replace
                if regex.search(inline[i].text) and not started:
                    print('work1')
                    found_runs.append((i, inline[i].text.find(regex.pattern), len(regex.pattern)))
                    text = inline[i].text.replace(regex.pattern, replace)
                    inline[i].text = text
                    print(text)
                    replace_done = True
                    found_all = True
                    
                elif regex.pattern[key_index] not in inline[i].text and not started:
                    # keep looking ...
                    print('smth')
                    # continue

                    start_index = inline[i].text.find(regex.pattern[key_index])
                    check_length = len(inline[i].text)
                    for text_index in range(start_index, check_length):
                        if inline[i].text[text_index] != regex.pattern[key_index]:
                            break
                    if key_index == 0:
                        started = True
                    chars_found = check_length - start_index
                    key_index += chars_found
                    found_runs.append((i, start_index, chars_found))
                    if key_index != len(regex.pattern):
                        continue
                    else:
                        found_all = True
                        break

                # case 2: search for partial text, find first run
                elif regex.pattern[key_index] not in inline[i].text[-1] in regex.pattern and not started:
                    print('work2')
                    start_index = inline[i].text.find(regex.pattern[key_index])
                    check_length = len(inline[i].text)
                    for text_index in range(start_index, check_length):
                        if inline[i].text[text_index] != regex.pattern[key_index]:
                            break
                    if key_index == 0:
                        started = True
                    chars_found = check_length - start_index
                    key_index += chars_found
                    found_runs.append((i, start_index, chars_found))
                    if key_index != len(regex.pattern):
                        continue
                    else:
                        found_all = True
                        break

                # case 3: search for partial text, find subsequent run
                elif regex.pattern[key_index] in inline[i].text and started and not found_all:
                    print('work3')
                    # check sequence
                    chars_found = 0
                    check_length = len(inline[i].text)
                    for text_index in range(0, check_length):
                        if inline[i].text[text_index] == regex.pattern[key_index]:
                            key_index += 1
                            chars_found += 1
                        else:
                            break
                    # no match so must be end
                    found_runs.append((i, 0, chars_found))
                    if key_index == len(regex.pattern):
                        found_all = True
                        break
              
            if found_all and not replace_done:
                print('all')
                for i, item in enumerate(found_runs):
                    index, start, length = [t for t in item]
                    if i == 0:
                        text = inline[index].text.replace(inline[index].text[start:start + length], regex.pattern)
                        inline[index].text = text
                    else:
                        text = inline[index].text.replace(inline[index].text[start:start + length], '')
                        inline[index].text = text
            # print(p.text)

            #     if regex.search(inline[i].text):
            #         text = regex.sub(replace, inline[i].text)
            #         inline[i].text = text
            #         print(text)




def repl_template(templ_name: str, file_name: str, temps: dict, info: dict):
    """

    :param templ_name: Имя файла шаблона
    :param file_name: Имя сохраняемого файла
    :param temps: заменяемые данные
    :param info: на что заменяем
    :return:
    """
    open_doc = Document(templ_name)

    list_tmps = [re.compile(f'<{key}>{value}>') for key, value in temps.items()]

    # txt_doc = d.process(templ_name)
    # patterns = re.compile(r"{(\w+)}\[(\w+.?\w+)]")
    # list_tmps = patterns.findall(str(txt_doc))

    dict_ = {key: value for key, value in zip(list(list_tmps), info.values())}
    print(dict_)

    for template, info in dict_.items():
        # print(str(template) + '->' + info)
        docx_replace_regex(open_doc, template, info)
    open_doc.save(file_name + '.docx')


def create_template_dict(docx_file):
    """
    Получение заменяемых переменных из выбранного шаблона
    :param docx_file: Шаблон документа
    :return: Словарь из заменяемого слова и его псевдонима
    """
    txt = d.process(docx_file)

    patterns = re.compile(r"<(\w+)>(\w+.?\w+)>")

    list_patterns = patterns.findall(str(txt))
    print(list_patterns)
    dict_ = {key: value for key, value in list_patterns}

    return dict_


def input_info(dict_tmps: dict):
    dict_info = {}
    for key, val in dict_tmps.items():
        if val == 'Дата':
            dict_info[key] = create_date()
        elif val == 'Номер договора':
            dict_info[key] = create_num(input(f"{val}: "))
        elif val == 'Инициалы':
            dict_info[key] = ''
        else:
            dict_info[key] = input(f"{val}: ")

    if 'initials' in dict_info.keys():
        last_name = dict_info.get('c_ln')
        name_letter = dict_info.get('c_n')[0]

        if dict_info.get('c_middn') == '':
            dict_info['initials'] = '.'.join([name_letter, last_name])
        else:
            middle_name = dict_info.get('c_middn')[0]
            dict_info['initials'] = '.'.join([name_letter, middle_name, last_name])

    return dict_info


def get_services(services) -> list:
    """
    Получение случайного списка имён

    :param services: длина списка
    :return: список с именами
    """
    services = SERVICES

    return services


def get_save_filename(open_tmpfile, inform_dict):
    filename = '.'.join((open_tmpfile.split('.')[:-1]))

    if 'num' in inform_dict.keys():
        name_save_doc = ''.join([filename, ' №', inform_dict['num']])
    else:
        name_save_doc = filename

    return name_save_doc


def main():
    list_temps = create_template_dict()
    list_info = input_info()

    tmp_doc = "Шаблон 1.docx"

    name_save_doc = ' '.join(['Договор', list_info[0]])
    repl_template(tmp_doc, name_save_doc, list_temps, list_info)


if __name__ == '__main__':
    open_docx_file = "Договор.docx"

    tmpls_dict = create_template_dict(docx_file=open_docx_file)
    info_dict = input_info(tmpls_dict)
    saving_filename = get_save_filename(open_docx_file, info_dict)

    repl_template(open_docx_file, saving_filename, tmpls_dict, info_dict)




